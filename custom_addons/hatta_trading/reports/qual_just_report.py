from odoo.addons.report_docx.report.report_docx import ReportDocx
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT



class QualJust(ReportDocx):

    def generate_docx_report(self, document, data, objs):

        sections = document.sections
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        p = document.add_paragraph()
        # run = document.add_paragraph().add_run()
        # font = run.font
        # font.size = Pt(16)

        p.style = document.styles['Normal']
        p.add_run('REQUEST FOR QUOTATION:').bold = True
        p.add_run().add_break()
        p.add_run('RFQ TITLE:' + '\t' + str(objs.reference)).add_break()
        p.add_run('RFQ No:' + '\t' + str(objs.customer_rfq)).add_break()
        p.add_run('OUR REF:' + '\t' + str(objs.coll_buyer_ref_no))

        # center aligned text
        p1 = document.add_paragraph()
        p1.add_run('APPENDIX 1').add_break()
        p1.add_run('COMMERCIAL BID SUBMISSION FORM')
        p1.alignment = 1

        #second para
        p2 = document.add_paragraph()
        p2.add_run('Date:' + '\t' + str(objs.creation_date)).add_break()
        p2.add_run('Name:' + '  ' + str(objs.partner_id.name)).add_break()
        p2.add_run(str(objs.partner_id.street) + '\n' + str(objs.partner_id.zip) + str(objs.partner_id.city)).add_break()
        p2.add_run('RFQ TITLE:' + '\t' + str(objs.reference)).add_break()
        p2.add_run('RFQ No.' + '\t' + str(objs.customer_rfq)).add_break()
        p2.add_run().add_break()
        p2.add_run('Dear Sir,').add_break()
        p2.add_run().add_break()

        p2.add_run('We have reviewed the Request for Quotation No.' + str(objs.customer_rfq) + ' ' + 'together with all the documentation   provided therewith or in connection thereto by the Company (RFQ).').add_break()
        p2.add_run().add_break()
        p2.add_run('We hereby offer to perform the Scope of Purchase (as defined in the RFQ) in conformity with the requirements stipulated in the RFQ, including the time period stipulated in the Scope of Purchase schedule in the RFQ, the fees provided by us in the Commercial Bid, and the Agreement, which we are ready to sign.').add_break()
        p2.add_run().add_break()
        p2.add_run('We confirm the validity of the fees provided in the Commercial Bid attached for the Scope of Purchase set out in the RFQ.').add_break()
        p2.add_run().add_break()
        p2.add_run('We confirm that we have satisfied ourselves as to the extent and nature of the Scope of Purchase to be provided by us and confirm that we have taken into account in our Bid all information issued by you during the Bid process.').add_break()
        p2.add_run().add_break()
        p2.add_run('We confirm that any and all qualifications to the RFQ have been either accepted by the Company in writing').add_break()
        p2.add_run().add_break()
        p2.add_run('Our Bid shall remain valid for [30] days from the Commercial Bid submission date and shall remain open for acceptance during this period.').add_break()
        p2.add_run().add_break()
        p2.add_run('We acknowledge and confirm that we shall be bound by the terms and conditions of the RFQ.').add_break()
        p2.add_run().add_break()
        p2.add_run('NOTE: Our offered prices do not include VAT charges, VAT charges will be invoiced separately at 	actuals.').add_break()
        p2.add_run().add_break()
        p2.add_run('Signed for an on behalf of [HATTA TRADING & SERVICES EST.]').add_break()
        p2.add_run().add_break()
        p2.add_run('By:______________________').add_break()
        p2.add_run().add_break()
        p2.add_run('Name:' + str(objs.user_id.name)).add_break()
        p2.add_run('Title').add_break()
        p2.add_run('Date:' + '\t' + str(objs.creation_date))

        # new page
        document.add_page_break()

        # page heading
        p3 = document.add_paragraph()
        p3.add_run('Technical Qualification').bold = True
        p3.alignment = 1

        p4 = document.add_paragraph()
        p4.add_run('REQUEST FOR QUOTATION').bold = True
        p4.add_run().add_break()
        p4.add_run('RFQ TITLE:' + '\t' + str(objs.reference)).add_break()
        p4.add_run('RFQ No.' + '\t' + str(objs.customer_rfq)).add_break()
        p4.add_run('OUR REF' + '\t' + str(objs.coll_buyer_ref_no))

        p5 = document.add_paragraph()
        p5.add_run('Appendix-2').bold = True
        p5.add_run().add_break()
        p5.add_run('Qualification Statement').bold = True
        p5.alignment = 1

        p6 = document.add_paragraph()
        p6.add_run('We hereby confirm that our Bid is in full and strict compliance with all the requirements stipulated in the RFQ, except for the following qualifications:')
        p6.add_run('Technical Qualifications :').bold = True

        # table
        table = document.add_table(rows=2+len(objs.selected_purchase_order_lines), cols=4)
        tr1_first_cell = table.cell(0, 0)
        table.style = 'TableGrid'
        table.autofit = True
        col1 = table.columns[0]
        col1.width = Cm(1)
        col2 = table.columns[1]
        col2.width = Cm(3)
        col3 = table.columns[2]
        col3.width = Cm(7)
        col4 = table.columns[3]
        col4.width = Cm(7)
        tr1_second_cell = table.cell(0, 3)
        merged_cell = tr1_first_cell.merge(tr1_second_cell)
        merged_cell.text = 'RFQ TITLE:' + '\t' + str(objs.reference) + '\n' + 'RFQ No:' + '\t' + str(objs.customer_rfq)
        tr2_cell1 = table.cell(1, 0)
        tr2_cell2 = table.cell(1, 1)
        tr2_cell3 = table.cell(1, 2)
        tr2_cell4 = table.cell(1, 3)
        tr2_cell1.text = 'No.'
        tr2_cell2.text = 'Ref No.'
        tr2_cell3.text = 'Qualification.'
        tr2_cell4.text = 'Justification'
        row = 2
        for item in objs.selected_purchase_order_lines:
            row_0 = table.cell(row, 0)
            row_0.text = str(item.serial_no)
            row_1 = table.cell(row, 1)
            row_1.text = str(item.product_id.default_code)
            row_2 = table.cell(row, 2)
            row_2.text = str(item.qualification)
            row_3 = table.cell(row, 3)
            row_3.text = str(item.justification)
            row += 1
        # para
        p7 = document.add_paragraph()
        p7.add_run().add_break()
        p7.add_run('NOTE:')

        p8 = document.add_paragraph()
        p8.add_run('Our offered prices do not include VAT charges, VAT charges will be invoiced separately at actuals.Kindly accept Partial Delivery Partial Payment Terms.')

        p9 = document.add_paragraph()

        p9.add_run('Duty Exemption Letter is required.').bold = True
        font = p9.add_run().font
        font.highlight_color
        font.highlight_color = WD_COLOR_INDEX.YELLOW

        p10 = document.add_paragraph()
        p10.add_run('We hereby confirm that any qualifications that are set out elsewhere in our Bid, but not listed in the above table, shall be deemed to be withdrawn and shall be considered null and void.')

        p10 = document.add_paragraph()
        p10.add_run('Signed by:')
        p10.add_run().add_break()
        p10.add_run().add_break()

        p10 = document.add_paragraph()
        p10.add_run('Name:' + str(objs.user_id.name)).add_break()
        p10.add_run('Title:').add_break()
        p10.add_run('Date:' + '\t' + str(objs.creation_date))


QualJust('report.hatta_trading.qual.docx', 'enquiry.details')